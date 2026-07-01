from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_campaign_docx(copy_data: dict, briefing_data: dict, output_path: str):
    """
    Genera un documento Word (.docx) professionale e formattato 
    contenente tutto il copy generato dall'agente Copywriter.
    """
    doc = Document()
    
    # Setup dei margini della pagina (A4 Standard)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Palette colori coerente con la UI (Noir & Cognac/Gold accent)
    COLOR_PRIMARY = RGBColor(19, 19, 21)     # #131315 (Noir)
    COLOR_ACCENT = RGBColor(195, 149, 90)    # #C3955A (Cognac)
    COLOR_MUTED = RGBColor(110, 110, 115)    # Gray
    
    # --- Intestazione Principale ---
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("FullForce.digital — Campaign Asset Sheet")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    title_p.paragraph_format.space_after = Pt(2)
    
    # Sottotitolo / Meta Informazioni
    meta_p = doc.add_paragraph()
    meta_text = f"Product: {briefing_data.get('product', '—')}  |  Season: {briefing_data.get('season', '—')}"
    meta_run = meta_p.add_run(meta_text)
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True
    meta_run.font.color.rgb = COLOR_MUTED
    meta_p.paragraph_format.space_after = Pt(24)
    
    # --- SEZIONE 1: COPY GENERATO ---
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Generated Copywriting Assets")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = COLOR_ACCENT
    h1.paragraph_format.space_after = Pt(12)
    
    # Tabella per mappare in modo pulito i campi del Copywriter
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asset Type'
    hdr_cells[1].text = 'Content'
    
    # Formattazione Header Tabella
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(11)
                
    # Estrazione sicura dei dati del copywriter
    copy_fields = [
        ("Top Label", copy_data.get("top_label", "—")),
        ("Headline", copy_data.get("headline", "—")),
        ("Subheadline", copy_data.get("subheadline", "—")),
        ("Trust Badges", ", ".join(copy_data.get("trust_badges", [])) if isinstance(copy_data.get("trust_badges"), list) else copy_data.get("trust_badges", "—"))
    ]
    
    for field_name, field_val in copy_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_val
        for cell in row_cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(10)
                    r.font.color.rgb = COLOR_PRIMARY
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # --- SEZIONE 2: DETTAGLI STRATEGICI ---
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Marketing Context & Guardrails")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = COLOR_ACCENT
    h2.paragraph_format.space_after = Pt(12)
    
    context_fields = [
        ("Target Audience", briefing_data.get("audience", "—")),
        ("Campaign Goal", briefing_data.get("goal", "—")),
        ("Tone of Voice", briefing_data.get("tone_of_voice", "—")),
    ]
    
    for label, val in context_fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_label = p.add_run(f"• {label}: ")
        r_label.font.bold = True
        r_label.font.name = 'Arial'
        r_label.font.size = Pt(10.5)
        r_label.font.color.rgb = COLOR_PRIMARY
        
        r_val = p.add_run(str(val))
        r_val.font.name = 'Arial'
        r_val.font.size = Pt(10.5)
        r_val.font.color.rgb = COLOR_PRIMARY

    # Salvataggio del file sul path specificato
    doc.save(output_path)