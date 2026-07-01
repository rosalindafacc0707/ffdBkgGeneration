"""
FastAPI Routes — Campaign API
"""
import logging
import base64
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from pathlib import Path
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import FileResponse, JSONResponse
from app.core.schemas import BriefingInput, BriefingJson, CampaignOutput, CopyOnlyOutput, VisualOnlyOutput, AssemblyOutput, ContentAssemblyInput
from app.core.config import settings
from app.agents.coordinator import run_campaign, run_copy, run_visual
from app.agents.brief_extractor import run_generate_brief_json
from app.mcp.workfront_mock import get_ready_briefings
from app.utils.azure_storage import upload_bytes_to_azure_blob


# Utils for the service of brief json extraction ---------
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/x-pdf",
    "application/acrobat",
}
MAX_FILE_SIZE_MB = 10
# ----------------------

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/v1", tags=["Campaign"])

@router.post(
    "/campaign/brief_insights_extraction",
    response_model=BriefingJson,
    summary="Generate valid JSON from campaign brief",
    description=(
        "Riceve un file PDF di briefing, analizza il contenuto ed estrae un JSON strutturato "
        "con tutti gli insight rilevanti per la generazione dell'asset finale della campagna. "
        "Il JSON restituito può essere usato direttamente come input di /campaign/generate."
    ),
)
async def generate_brief_json(
    file: UploadFile = File(
        ...,
        description="File PDF del briefing di campagna (max 10 MB)",
    )
) -> BriefingJson:
    # Validazione tipo file
    content_type = file.content_type or ""
    filename = file.filename or "unknown.pdf"

    if not filename.lower().endswith(".pdf") and content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=415,
            detail=(
                f"Formato file non supportato: '{content_type}'. "
                "Carica un file PDF (.pdf)."
            ),
        )

    # Leggi i bytes del file
    pdf_bytes = await file.read()

    # Validazione dimensione
    size_mb = len(pdf_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(
            status_code=413,
            detail=f"File troppo grande: {size_mb:.1f} MB. Dimensione massima: {MAX_FILE_SIZE_MB} MB.",
        )

    if len(pdf_bytes) == 0:
        raise HTTPException(status_code=400, detail="Il file caricato è vuoto.")

    logger.info(
        "API brief_insights_extraction: ricevuto '%s' (%.1f MB)", filename, size_mb
    )

    try:
        result = await run_generate_brief_json(pdf_bytes, filename)
        logger.info("API brief_insights_extraction: completato con status OK")
        return result
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        logger.error("API brief_insights_extraction error: %s", e)
        raise HTTPException(status_code=500, detail=f"Errore estrazione brief: {e}")
    

@router.post(
    "/campaign/generate_copy_and_background",
    response_model=CampaignOutput,
    summary="Campaign copy & image background generation",
    description=(
        "Receives a structured JSON briefing and runs the full agentic system. "
        "This returns both the advertising copy and the generated background image."
    ),
)
async def generate_campaign(briefing: BriefingInput) -> CampaignOutput:
    try:
        logger.info("API: richiesta campagna per prodotto '%s'", briefing.product)
        result = await run_campaign(briefing)
        return result
    except Exception as e:
        logger.error("API error: %s", e)
        raise HTTPException(status_code=500, detail=f"Errore generazione campagna: {e}")


@router.post(
    "/campaign/generate_copy",
    summary="Generate campaign copy only",
    response_model=CopyOnlyOutput,
    description=(
        "Genera solo la parte testuale della campagna (top_label/headline/subheadline/trust_badges). "
        "This endpoint returns copy immediately while image generation can run later."
    ),
)
async def generate_copy(briefing: BriefingInput):
    try:
        logger.info("API: richiesta copy per prodotto '%s'", briefing.product)
        copy_result = await run_copy(briefing)
        return {"copy": copy_result, "status": "copy_generated"}
    except Exception as e:
        logger.error("API error: %s", e)
        raise HTTPException(status_code=500, detail=f"Errore generazione copy: {e}")


@router.post(
    "/campaign/generate_background",
    summary="Generate campaign background image only",
    response_model=VisualOnlyOutput,
    description=(
        "Generates only the visual background image from an existing briefing. "
        "Use this after copy has been returned to generate the image in a separate request."
    ),
)
async def generate_background(briefing: BriefingInput):
    try:
        logger.info("API: richiesta immagine per prodotto '%s'", briefing.product)
        visual_result = await run_visual(briefing)
        return {"visual": visual_result, "status": "visual_generated"}
    except Exception as e:
        logger.error("API error: %s", e)
        raise HTTPException(status_code=500, detail=f"Errore generazione immagine: {e}")


@router.get(
    "/campaign/image/{filename}",
    summary="Scarica immagine background generata",
    description="Restituisce il file PNG del background generato.",
)
async def get_image(filename: str):
    from app.core.config import settings
    filepath = Path(settings.image_output_dir) / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Immagine non trovata")
    return FileResponse(str(filepath), media_type="image/png", filename=filename)



@router.get("/health", summary="Health check")
async def health():
    return {"status": "ok", "service": "FullForce Ad Generator"}


# ─────────────────────────────────────────────────────────────────────────
# MOCK — Workfront Integration
# Endpoint REST diretto, comodo per test rapidi da Swagger/curl senza
# passare dal protocollo generico /mcp/call. Internamente richiama lo
# stesso mock usato dal tool MCP `get_ready_briefings` (vedi app/mcp/server.py).
# Quando sara' disponibile il vero Workfront MCP server, sostituire
# l'implementazione in app/mcp/workfront_mock.py mantenendo invariata
# questa route.
# ─────────────────────────────────────────────────────────────────────────
@router.get(
    "/workfront/ready_briefings",
    summary="[MOCK] Recupera briefing in stato Ready da Workfront",
    description=(
        "Simula la chiamata al Workfront MCP server per recuperare i briefing "
        "di campagna attualmente in stato 'Ready'. Ogni elemento restituito "
        "contiene un `briefing_payload` direttamente compatibile con "
        "/campaign/brief_insights_extraction e /campaign/generate_copy_and_background. "
        "Nessuna chiamata reale viene effettuata verso Workfront — dati mock per sviluppo/demo."
    ),
)
async def get_ready_briefings_mock(limit: int = 5, project_id: str | None = None):
    try:
        logger.info("API: richiesta mock Workfront ready briefings (limit=%d)", limit)
        result = await get_ready_briefings(limit=limit, project_id=project_id)
        return JSONResponse(content=result)
    except Exception as e:
        logger.error("API error (workfront mock): %s", e)
        raise HTTPException(status_code=500, detail=f"Errore mock Workfront: {e}")

@router.post(
    "/campaign/assemble_content",
    response_model=AssemblyOutput,
    summary="Upload assets to Azure Blob Storage",
    description="Salva la copia generata in formato .docx e l'immagine corrente in formato .png in Azure Blob Storage.",
)
async def assemble_content(payload: ContentAssemblyInput):
    timestamp = int(time.time())
    product_slug = (payload.product or "product").lower().replace(" ", "_")[:20]

    saved_files = []

    # 1. SALVATAGGIO IMMAGINE .PNG
    if payload.image_base64:
        try:
            img_data = base64.b64decode(payload.image_base64)
            blob_name = f"images/bg_{product_slug}_{timestamp}.png"
            image_url = upload_bytes_to_azure_blob(img_data, blob_name, content_type="image/png")
            saved_files.append(image_url)
            logger.info(" [ASSEMBLY] Immagine caricata su Azure Blob: %s", image_url)
        except Exception as e:
            logger.error(" [ASSEMBLY] Errore salvataggio immagine: %s", e)
            raise HTTPException(status_code=500, detail=f"Errore caricamento PNG su Azure: {e}")

    # 2. SALVATAGGIO COPY .DOCX IN STRUTTURA PROFESSIONALE
    try:
        doc = Document()
        
        # Setup Margini standard aziendali
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Palette Colore Professionale basata sui token della UI
        COLOR_PRIMARY = RGBColor(19, 19, 21)     # Noir (#131315)
        COLOR_ACCENT = RGBColor(195, 149, 90)    # Cognac (#C3955A)
        COLOR_MUTED = RGBColor(122, 120, 118)    # Text-dim (#7a7876)

        # Intestazione Principale
        p_title = doc.add_paragraph()
        r_title = p_title.add_run("FullForce.digital — Campaign Asset Sheet")
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(20)
        r_title.font.bold = True
        r_title.font.color.rgb = COLOR_PRIMARY
        p_title.paragraph_format.space_after = Pt(2)

        # Meta-dati della campagna
        p_meta = doc.add_paragraph()
        meta_text = f"Product: {payload.product}  |  Season: {payload.season}"
        if payload.brand:
            meta_text += f"  |  Brand: {payload.brand}"
        r_meta = p_meta.add_run(meta_text)
        r_meta.font.name = 'Arial'
        r_meta.font.size = Pt(10)
        r_meta.font.italic = True
        r_meta.font.color.rgb = COLOR_MUTED
        p_meta.paragraph_format.space_after = Pt(24)

        # --- SEZIONE COPYWRITING ASSETS ---
        h1 = doc.add_paragraph()
        h1_run = h2_run = h1.add_run("1. Final Copywriting Components")
        h1_run.font.name = 'Arial'
        h1_run.font.size = Pt(14)
        h1_run.font.bold = True
        h1_run.font.color.rgb = COLOR_ACCENT
        h1.paragraph_format.space_after = Pt(12)

        # Tabella pulita per il layout di copy
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Component'
        hdr_cells[1].text = 'Content Text'
        
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.name = 'Arial'
                    r.font.size = Pt(11)

        # Mapping dei dati estratti dal frontend (con fallback in caso di stringa vuota o trattino)
        copy_data = [
            ("Top Label", payload.custom_top_label),
            ("Headline", payload.custom_headline),
            ("Description / Subheadline", payload.custom_subheadline)
        ]

        for label, val in copy_data:
            clean_val = val.strip() if val else "—"
            if clean_val == "" or clean_val == "—":
                clean_val = "—"
                
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = clean_val
            
            for cell in row_cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(10.5)
                        r.font.color.rgb = COLOR_PRIMARY

        doc.add_paragraph().paragraph_format.space_after = Pt(24)

        # --- SEZIONE PROMPT DI GENERAZIONE ---
        if payload.image_prompt:
          h2 = doc.add_paragraph()
          h2_run = h2.add_run("2. Applied Image Prompt Context")
          h2_run.font.name = 'Arial'
          h2_run.font.size = Pt(14)
          h2_run.font.bold = True
          h2_run.font.color.rgb = COLOR_ACCENT
          h2.paragraph_format.space_after = Pt(8)

          p_prompt = doc.add_paragraph()
          r_prompt = p_prompt.add_run(payload.image_prompt)
          r_prompt.font.name = 'Courier New'
          r_prompt.font.size = Pt(9.5)
          r_prompt.font.color.rgb = COLOR_PRIMARY
          p_prompt.paragraph_format.space_after = Pt(12)

        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        doc_blob_name = f"documents/copy_{product_slug}_{timestamp}.docx"
        doc_url = upload_bytes_to_azure_blob(doc_bytes.getvalue(), doc_blob_name, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        saved_files.append(doc_url)
        logger.info(" [ASSEMBLY] Documento DOCX caricato su Azure Blob: %s", doc_url)

    except Exception as e:
        logger.error(" [ASSEMBLY] Errore salvataggio documento DOCX: %s", e)
        raise HTTPException(status_code=500, detail=f"Errore caricamento DOCX su Azure: {e}")
        
    return {
        "status": "success",
        "saved_files": saved_files
    }